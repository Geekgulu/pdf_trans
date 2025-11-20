import os
from typing import List, Tuple
from dotenv import load_dotenv
import requests
import pdfplumber
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
import streamlit as st
import tempfile
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import urllib.request
import zipfile
import shutil

class PDFTranslator:
    def __init__(self):
        load_dotenv()
        self.api_key = os.getenv('DEEPSEEK_API_KEY')
        self.api_url = 'https://api.siliconflow.cn/v1/chat/completions'
        self.headers = {
            'Authorization': f'Bearer {self.api_key}',
            'Content-Type': 'application/json'
        }
        # 确保字体目录存在
        self.fonts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'fonts')
        if not os.path.exists(self.fonts_dir):
            os.makedirs(self.fonts_dir)

    def extract_text_from_pdf(self, pdf_path: str) -> List[Tuple[int, dict]]:
        """从PDF文件中提取文本、表格和图片，按段落返回内容"""
        content_by_page = []
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total_pages = len(pdf.pages)
                # 创建文本提取进度条
                extract_progress = st.progress(0)
                extract_status = st.empty()
                
                for page_num, page in enumerate(pdf.pages, start=1):
                    extract_status.text(f"正在提取第 {page_num}/{total_pages} 页...")
                    extract_progress.progress(page_num / total_pages)
                    
                    # 获取页面文本
                    text = page.extract_text() or ""
                    
                    # 使用更精确的段落分割方法
                    paragraphs = []
                    # 按两个或更多换行符分割段落
                    raw_paragraphs = text.split('\n\n')
                    
                    for para in raw_paragraphs:
                        para = para.strip()
                        if para:
                            paragraphs.append({
                                'text': para,
                                'bbox': None
                            })

                    page_content = {
                        "paragraphs": paragraphs,
                        "tables": page.extract_tables(),
                        "images": page.images
                    }
                    content_by_page.append((page_num, page_content))
                
                extract_status.text("文本提取完成！")
                extract_progress.progress(1.0)
                return content_by_page
                
        except Exception as e:
            raise Exception(f'PDF文件读取失败：{str(e)}')

    def translate_text(self, text: str, target_lang: str) -> str:
        """调用Deepseek API翻译文本"""
        if not isinstance(text, str):
            raise ValueError("输入的文本必须是字符串类型")
        if not text.strip():
            return ""  # 如果文本为空，直接返回空字符串

        system_prompt = f"""你是一个专业的翻译助手。请严格按照以下要求翻译文本：
1. 只输出翻译后的内容，不要添加任何解释、注释或说明
2. 保持原文的格式和段落结构
3. 翻译成{target_lang}
4. 不要输出"翻译如下"、"以下是翻译"、“原文”、“译文”等提示性文字
5. 不要添加任何括号内的解释或补充说明
6. 直接输出翻译结果，不要有任何前缀或后缀"""
        try:
            response = requests.post(
                self.api_url,
                headers=self.headers,
                json={
                    'model': 'deepseek-ai/DeepSeek-V3',
                    'messages': [
                        {'role': 'system', 'content': system_prompt},
                        {'role': 'user', 'content': text}
                    ],
                    'max_tokens': 1000,  # 限制最大 token 数，避免超出限制
                    'temperature': 0.3
                }
            )
            response.raise_for_status()
            result = response.json()
            if "choices" in result and len(result["choices"]) > 0:
                return result["choices"][0]["message"]["content"].strip()
            else:
                raise Exception("翻译API返回的结果格式不正确")
        except Exception as e:
            raise Exception(f'翻译请求失败：{str(e)}')

    def create_translated_pdf(self, original_pdf: str, original_texts: List[Tuple[int, dict]], 
                        translated_texts: List[Tuple[int, dict]], output_path: str, show_comparison: bool = True):
        """创建翻译后的PDF文件，支持原文译文对照"""
        try:
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            
            # 创建PDF文档
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # 注册多语言字体
            font_name = self._register_fonts()
            
            # 创建样式
            styles = getSampleStyleSheet()
            original_style = ParagraphStyle(
                'OriginalText',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
                leading=16,
                spaceAfter=6,
                alignment=0,  # 左对齐
                textColor='#666666'
            )
            translated_style = ParagraphStyle(
                'TranslatedText',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=12,
                leading=18,
                spaceAfter=12,
                alignment=0  # 左对齐
            )
            
            # 构建文档内容
            story = []
            
            for (orig_page, orig_content), (trans_page, trans_content) in zip(original_texts, translated_texts):
                # 处理段落
                for orig_para, trans_para in zip(orig_content["paragraphs"], trans_content["paragraphs"]):
                    if show_comparison:
                        # 添加原文
                        if orig_para["text"].strip():
                            story.append(Paragraph(orig_para["text"].strip(), original_style))
                        # 添加译文
                        if trans_para["text"].strip():
                            story.append(Paragraph(trans_para["text"].strip(), translated_style))
                    else:
                        # 仅显示译文
                        if trans_para["text"].strip():
                            story.append(Paragraph(trans_para["text"].strip(), translated_style))
                
                story.append(Spacer(1, 12))  # 段落间距

                # 添加表格
                for table in orig_content["tables"]:
                    table_data = [[Paragraph(cell or "", original_style) for cell in row] for row in table]
                    table_obj = Table(table_data)
                    table_obj.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('FONTNAME', (0, 0), (-1, -1), font_name),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ]))
                    story.append(table_obj)
                    story.append(Spacer(1, 12))
                
                # 添加图片
                for image in orig_content["images"]:
                    img = Image(image["src"], width=inch, height=inch)
                    story.append(img)
                    story.append(Spacer(1, 12))
        
            # 生成PDF
            doc.build(story)
    
        except Exception as e:
            raise Exception(f'PDF文件创建失败：{str(e)}')
    
    def _register_fonts(self):
        """注册多语言字体，如果本地没有则自动下载"""
        # 字体文件路径
        font_paths = {
            'ArialUnicode': 'C:/Windows/Fonts/ARIALUNI.TTF',
            'NotoSans': 'C:/Windows/Fonts/NotoSans-Regular.ttf',
            'SimSun': 'C:/Windows/Fonts/simsun.ttc',
            'SimHei': 'C:/Windows/Fonts/simhei.ttf',
            'SourceHanSans': os.path.join(self.fonts_dir, 'SourceHanSansSC-Regular.otf')
        }
        
        # 尝试注册已有字体
        for font_name, font_path in font_paths.items():
            try:
                if os.path.exists(font_path):
                    from reportlab.pdfbase import pdfmetrics
                    from reportlab.pdfbase.ttfonts import TTFont
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    return font_name
            except Exception as e:
                st.warning(f"注册字体 {font_name} 失败: {str(e)}")
                continue
        
        # 如果没有可用字体，下载并安装思源黑体
        try:
            source_han_path = font_paths['SourceHanSans']
            if not os.path.exists(source_han_path):
                st.info("正在下载思源黑体字体，请稍候...")
                # 下载思源黑体
                font_url = "https://github.com/adobe-fonts/source-han-sans/releases/download/2.004R/SourceHanSansSC.zip"
                zip_path = os.path.join(self.fonts_dir, "SourceHanSansSC.zip")
                
                urllib.request.urlretrieve(font_url, zip_path)
                
                # 解压字体文件
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(self.fonts_dir)
                
                # 删除zip文件
                os.remove(zip_path)
                
                st.success("思源黑体字体下载完成！")
            
            # 注册思源黑体
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            pdfmetrics.registerFont(TTFont('SourceHanSans', source_han_path))
            return 'SourceHanSans'
        except Exception as e:
            st.error(f"下载安装思源黑体失败: {str(e)}")
            # 最后的备选方案
            return 'Helvetica'

    # 添加Word文档处理方法
    def extract_text_from_docx(self, docx_path: str) -> List[Tuple[int, dict]]:
        """从Word文档中提取文本，按段落返回内容"""
        content_by_page = []
        try:
            doc = Document(docx_path)
            
            # 创建文本提取进度条
            extract_progress = st.progress(0)
            extract_status = st.empty()
            extract_status.text("正在提取Word文档内容...")
            
            # 将文档内容按段落组织
            paragraphs = []
            tables = []
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text.strip(),
                        'bbox': None
                    })
            
            # 提取表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
            
            # 由于Word文档没有明确的页面概念，我们将整个文档视为一页
            page_content = {
                "paragraphs": paragraphs,
                "tables": tables,
                "images": []  # Word文档中的图片处理较复杂，暂不支持
            }
            content_by_page.append((1, page_content))
            
            extract_status.text("文本提取完成！")
            extract_progress.progress(1.0)
            return content_by_page
                
        except Exception as e:
            raise Exception(f'Word文档读取失败：{str(e)}')

    def create_interleaved_docx(self, input_file: str, translated_texts: List[Tuple[int, dict]], output_path: str):
        """创建交错的Word文档：按段落原文-译文交替显示，不包含任何提示性标题"""
        try:
            doc = Document()

            # 默认样式设置为宋体，并设置 East Asia 字体映射，避免中文乱码
            style = doc.styles['Normal']
            style.font.name = 'SimSun'
            style.font.size = Pt(12)
            # 设置样式的 East Asia 字体（兼容中文）
            if hasattr(style, 'element') and style.element is not None:
                rFonts = style.element.rPr.rFonts
                rFonts.set(qn('w:eastAsia'), 'SimSun')

            # 读取原始文档内容
            original_doc = Document(input_file)

            original_paragraphs = [p.text.strip() for p in original_doc.paragraphs if p.text.strip()]
            original_tables = []
            for table in original_doc.tables:
                t = []
                for row in table.rows:
                    t.append([cell.text for cell in row.cells])
                if t:
                    original_tables.append(t)

            # 译文数据（docx视为单页）
            translated_page = translated_texts[0][1] if translated_texts else {"paragraphs": [], "tables": []}
            translated_paragraphs = [p.get("text", "").strip() for p in translated_page.get("paragraphs", [])]
            translated_tables = translated_page.get("tables", [])

            # 段落：原文与译文交替，无标题、无分隔线
            n = max(len(original_paragraphs), len(translated_paragraphs))
            for i in range(n):
                if i < len(original_paragraphs):
                    p = doc.add_paragraph(original_paragraphs[i])
                    p.paragraph_format.first_line_indent = Pt(24)
                    for run in p.runs:
                        run.font.name = 'SimSun'
                        # 设置 East Asia 字体以防乱码
                        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

                if i < len(translated_paragraphs):
                    p = doc.add_paragraph(translated_paragraphs[i])
                    p.paragraph_format.first_line_indent = Pt(24)
                    for run in p.runs:
                        run.font.name = 'SimSun'
                        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            # 表格：原表格后紧接译文表格，不加任何标题
            m = max(len(original_tables), len(translated_tables))
            for i in range(m):
                if i < len(original_tables):
                    table_data = original_tables[i]
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(table_data):
                        for c_idx, cell in enumerate(row):
                            table.cell(r_idx, c_idx).text = cell or ""
                            for paragraph in table.cell(r_idx, c_idx).paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'SimSun'
                                    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

                if i < len(translated_tables):
                    trans_data = translated_tables[i]
                    table = doc.add_table(rows=len(trans_data), cols=len(trans_data[0]))
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(trans_data):
                        for c_idx, cell in enumerate(row):
                            table.cell(r_idx, c_idx).text = cell or ""
                            for paragraph in table.cell(r_idx, c_idx).paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'SimSun'
                                    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            doc.save(output_path)

        except Exception as e:
            raise Exception(f'Word文档创建失败：{str(e)}')

    def create_translated_docx(self, translated_texts: List[Tuple[int, dict]], output_path: str, show_comparison: bool = True):
        """创建仅译文的Word文档（无任何提示性标题），按段落排版"""
        try:
            doc = Document()

            # 默认样式：宋体 + East Asia 映射（防乱码）
            style = doc.styles['Normal']
            style.font.name = 'SimSun'
            style.font.size = Pt(12)
            if hasattr(style, 'element') and style.element is not None:
                style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            # 遍历所有译文页（docx通常只有一页结构）
            for _, page_content in translated_texts:
                for para in page_content.get("paragraphs", []):
                    text = (para.get("text") or "").strip()
                    if text:
                        p = doc.add_paragraph(text)
                        p.paragraph_format.first_line_indent = Pt(24)
                        for run in p.runs:
                            run.font.name = 'SimSun'
                            if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

                for table_data in page_content.get("tables", []):
                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Table Grid'
                        for r_idx, row in enumerate(table_data):
                            for c_idx, cell in enumerate(row):
                                table.cell(r_idx, c_idx).text = cell or ""
                                for paragraph in table.cell(r_idx, c_idx).paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'SimSun'
                                        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            doc.save(output_path)

        except Exception as e:
            raise Exception(f'Word文档创建失败：{str(e)}')

    def _create_translation_pages(self, translated_texts: List[Tuple[int, dict]], output_path: str):
        """创建译文页面（无提示性标题），每页按段落排版"""
        try:
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            from reportlab.lib import colors

            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )

            font_name = self._register_fonts()
            styles = getSampleStyleSheet()

            translated_style = ParagraphStyle(
                'TranslatedText',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=12,
                leading=18,
                spaceBefore=6,
                spaceAfter=6,
                alignment=0  # 左对齐
            )

            story = []

            for i, (page_num, page_content) in enumerate(translated_texts):
                if i > 0:
                    story.append(PageBreak())

                # 仅添加译文段落（不再添加“第 X 页译文”等标题）
                paragraphs = page_content.get("paragraphs", [])
                for para in paragraphs:
                    text = (para.get("text") or "").strip()
                    if text:
                        story.append(Paragraph(text, translated_style))
                        story.append(Spacer(1, 8))

                # 表格（不添加“表格数据”、“表格 X”等提示）
                for table_data in page_content.get("tables", []):
                    if table_data:
                        table = Table(table_data, repeatRows=0)
                        table.setStyle(TableStyle([
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                            ('FONTNAME', (0, 0), (-1, -1), font_name),
                            ('FONTSIZE', (0, 0), (-1, -1), 10),
                            ('PADDING', (0, 0), (-1, -1), 6),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ]))
                        story.append(table)
                        story.append(Spacer(1, 12))

            doc.build(story)

        except Exception as e:
            raise Exception(f'译文页面创建失败：{str(e)}')
    
    def create_interleaved_pdf(self, original_pdf: str, translated_texts: List[Tuple[int, dict]], output_path: str):
        """创建交错的PDF文件，原文页面和译文页面交替出现"""
        temp_trans_path = None
    
        try:
            # 使用临时文件上下文管理器
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_trans_path = temp_file.name
    
            # 创建译文页面
            self._create_translation_pages(translated_texts, temp_trans_path)
    
            # 检查临时文件是否成功创建
            if not os.path.exists(temp_trans_path):
                raise Exception("临时文件创建失败")
    
            # 创建最终的PDF文件
            writer = PdfWriter()
    
            # 打开原始PDF和译文PDF
            with open(original_pdf, 'rb') as orig_file, open(temp_trans_path, 'rb') as trans_file:
                # 一次性读取两个PDF文件
                orig_reader = PdfReader(orig_file)
                trans_reader = PdfReader(trans_file)
    
                # 计算总页数
                total_orig_pages = len(orig_reader.pages)
                total_trans_pages = len(trans_reader.pages)
    
                # 添加所有页面
                for i in range(max(total_orig_pages, total_trans_pages)):
                    # 添加原文页面
                    if i < total_orig_pages:
                        writer.add_page(orig_reader.pages[i])
    
                    # 添加译文页面
                    if i < total_trans_pages:
                        writer.add_page(trans_reader.pages[i])
    
            # 写入最终的PDF文件
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
    
        except Exception as e:
            raise Exception(f'PDF交错合并失败：{str(e)}')
    
        finally:
            # 删除临时文件
            if temp_trans_path and os.path.exists(temp_trans_path):
                try:
                    os.unlink(temp_trans_path)
                except:
                    pass

    def translate_pdf(self, input_file: str, output_file: str, target_language: str, show_comparison: bool = True):
        """翻译PDF文档"""
        try:
            # 提取PDF文档内容
            extracted_texts = self.extract_text_from_pdf(input_file)
            
            # 创建翻译进度条
            total_paragraphs = sum(len(page_content["paragraphs"]) for _, page_content in extracted_texts)
            translate_progress = st.progress(0)
            translate_status = st.empty()
            
            # 翻译所有段落
            translated_texts = []
            for page_num, page_content in extracted_texts:
                translated_paragraphs = []
                translated_tables = []
                
                # 翻译段落
                for i, para in enumerate(page_content["paragraphs"]):
                    translate_status.text(f"正在翻译第 {i+1}/{total_paragraphs} 个段落...")
                    translate_progress.progress((i+1) / total_paragraphs)
                    
                    if para["text"].strip():
                        translated_text = self.translate_text(para["text"], target_language)
                        translated_paragraphs.append({
                            "text": translated_text,
                            "bbox": para["bbox"]
                        })
                    else:
                        translated_paragraphs.append(para)
                
                # 翻译表格
                for table in page_content["tables"]:
                    translated_table = []
                    for row in table:
                        translated_row = []
                        for cell in row:
                            if cell and cell.strip():
                                translated_cell = self.translate_text(cell, target_language)
                                translated_row.append(translated_cell)
                            else:
                                translated_row.append(cell)
                        translated_table.append(translated_row)
                    translated_tables.append(translated_table)
                # 删除这行重复代码: translated_tables.append(translated_table)
                
                # 组合翻译后的内容
                translated_page_content = {
                    "paragraphs": translated_paragraphs,
                    "tables": translated_tables,
                    "images": page_content["images"]
                }
                translated_texts.append((page_num, translated_page_content))
            
            translate_status.text("翻译完成！正在生成PDF文档...")
            
            # 创建翻译后的PDF文档
            if show_comparison:
                self.create_interleaved_pdf(input_file, translated_texts, output_file)
            else:
                self._create_translation_pages(translated_texts, output_file)
            
            translate_status.text("PDF文档生成完成！")
            translate_progress.progress(1.0)
            
            return True
                
        except Exception as e:
            st.error(f"翻译失败: {str(e)}")
            raise Exception(f'文档翻译失败：{str(e)}')

    def translate_document(self, input_file: str, output_file: str, target_language: str, show_comparison: bool = True, file_type: str = 'pdf'):
        """翻译文档（支持PDF和Word文档）"""
        try:
            if file_type.lower() == 'pdf':
                return self.translate_pdf(input_file, output_file, target_language, show_comparison)
            elif file_type.lower() == 'docx':
                # 提取Word文档内容
                extracted_texts = self.extract_text_from_docx(input_file)
                
                # 创建翻译进度条
                total_paragraphs = sum(len(page_content["paragraphs"]) for _, page_content in extracted_texts)
                translate_progress = st.progress(0)
                translate_status = st.empty()
                
                # 翻译所有段落
                translated_texts = []
                for page_num, page_content in extracted_texts:
                    translated_paragraphs = []
                    translated_tables = []
                    
                    # 翻译段落
                    for i, para in enumerate(page_content["paragraphs"]):
                        translate_status.text(f"正在翻译第 {i+1}/{total_paragraphs} 个段落...")
                        translate_progress.progress((i+1) / total_paragraphs)
                        
                        if para["text"].strip():
                            translated_text = self.translate_text(para["text"], target_language)
                            translated_paragraphs.append({
                                "text": translated_text,
                                "bbox": para["bbox"]
                            })
                        else:
                            translated_paragraphs.append(para)
                    
                    # 翻译表格
                    for table in page_content["tables"]:
                        translated_table = []
                        for row in table:
                            translated_row = []
                            for cell in row:
                                if cell and cell.strip():
                                    translated_cell = self.translate_text(cell, target_language)
                                    translated_row.append(translated_cell)
                                else:
                                    translated_row.append(cell)
                            translated_table.append(translated_row)
                        translated_tables.append(translated_table)
                    
                    # 组合翻译后的内容
                    translated_page_content = {
                        "paragraphs": translated_paragraphs,
                        "tables": translated_tables,
                        "images": page_content["images"]
                    }
                    translated_texts.append((page_num, translated_page_content))
                
                translate_status.text("翻译完成！正在生成Word文档...")
                
                # 创建翻译后的Word文档
                if show_comparison:
                    self.create_interleaved_docx(input_file, translated_texts, output_file)
                else:
                    self.create_translated_docx(translated_texts, output_file)
                
                translate_status.text("Word文档生成完成！")
                translate_progress.progress(1.0)
                
                return True
            else:
                raise Exception(f'不支持的文件类型：{file_type}')
                
        except Exception as e:
            st.error(f"翻译失败: {str(e)}")
            raise Exception(f'文档翻译失败：{str(e)}')

def main():
    # 使用示例
    translator = PDFTranslator()
    input_file = 'input.pdf'  # 输入文件路径
    output_file = 'output.pdf'  # 输出文件路径
    target_language = '中文'  # 目标语言：'中文' 或 'English'
    
    translator.translate_pdf(input_file, output_file, target_language)

if __name__ == '__main__':
    main()